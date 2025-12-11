# Converts PDF, Word (.docx), and TXT files to Markdown format

import fitz  # PyMuPDF
from docx import Document
from io import BytesIO
import re


def pdf_to_markdown(content: bytes) -> str:
    doc = fitz.open(stream=content, filetype="pdf")
    markdown_parts = []
    
    for page_num, page in enumerate(doc):
        page_md = []
        
        # Try to find tables first
        try:
            tables = page.find_tables()
            table_rects = [fitz.Rect(t.bbox) for t in tables.tables] if tables.tables else []
        except:
            table_rects = []
        
        # Get text blocks
        blocks = page.get_text("dict")["blocks"]
        
        for block in blocks:
            if block.get("type") == 0:  # Text block
                block_rect = fitz.Rect(block["bbox"])
                
                # Skip if inside table (we'll handle tables separately)
                if any(block_rect.intersects(tr) for tr in table_rects):
                    continue
                
                for line in block.get("lines", []):
                    line_text = ""
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue
                        
                        font_size = span.get("size", 12)
                        flags = span.get("flags", 0)
                        is_bold = flags & 2 ** 4  # Bold flag
                        
                        # Detect headers based on font size
                        if font_size >= 20:
                            text = f"# {text}"
                        elif font_size >= 16:
                            text = f"## {text}"
                        elif font_size >= 14:
                            text = f"### {text}"
                        elif is_bold and len(text) < 100:
                            text = f"**{text}**"
                        
                        line_text += text + " "
                    
                    if line_text.strip():
                        page_md.append(line_text.strip())
        
        # Handle tables
        try:
            tables = page.find_tables()
            if tables.tables:
                for table in tables.tables:
                    table_data = table.extract()
                    if table_data and len(table_data) > 0:
                        page_md.append("")
                        # Header row
                        headers = table_data[0]
                        header_row = "| " + " | ".join([str(h) if h else "" for h in headers]) + " |"
                        separator = "| " + " | ".join(["---" for _ in headers]) + " |"
                        page_md.append(header_row)
                        page_md.append(separator)
                        
                        # Data rows
                        for row in table_data[1:]:
                            row_text = "| " + " | ".join([str(cell) if cell else "" for cell in row]) + " |"
                            page_md.append(row_text)
                        page_md.append("")
        except:
            pass
        
        if page_md:
            markdown_parts.append(f"\n---\n*Page {page_num + 1}*\n\n" + "\n".join(page_md))
    
    doc.close()
    return "\n\n".join(markdown_parts)


def docx_to_markdown(content: bytes) -> str:
    doc = Document(BytesIO(content))
    markdown_parts = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            markdown_parts.append("")
            continue
        
        style_name = para.style.name.lower() if para.style else ""
        
        # Detect headings
        if "heading 1" in style_name:
            markdown_parts.append(f"# {text}")
        elif "heading 2" in style_name:
            markdown_parts.append(f"## {text}")
        elif "heading 3" in style_name:
            markdown_parts.append(f"### {text}")
        elif "heading 4" in style_name:
            markdown_parts.append(f"#### {text}")
        elif "title" in style_name:
            markdown_parts.append(f"# {text}")
        elif "list" in style_name:
            markdown_parts.append(f"- {text}")
        else:
            # Check for inline formatting
            formatted_text = ""
            for run in para.runs:
                run_text = run.text
                if run.bold and run.italic:
                    run_text = f"***{run_text}***"
                elif run.bold:
                    run_text = f"**{run_text}**"
                elif run.italic:
                    run_text = f"*{run_text}*"
                formatted_text += run_text
            
            markdown_parts.append(formatted_text if formatted_text else text)
    
    # Handle tables
    for table in doc.tables:
        markdown_parts.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            row_text = "| " + " | ".join(cells) + " |"
            markdown_parts.append(row_text)
            
            # Add separator after header row
            if i == 0:
                separator = "| " + " | ".join(["---" for _ in cells]) + " |"
                markdown_parts.append(separator)
        markdown_parts.append("")
    
    return "\n".join(markdown_parts)


def txt_to_markdown(content: bytes) -> str:

    text = content.decode("utf-8", errors="ignore")
    
    lines = text.split("\n")
    markdown_parts = []
    
    for line in lines:
        line = line.strip()
        
        # Detect potential headers (ALL CAPS lines, short lines)
        if line and line.isupper() and len(line) < 80:
            markdown_parts.append(f"## {line.title()}")
        # Detect bullet points
        elif line.startswith(("- ", "• ", "* ", "· ")):
            markdown_parts.append(f"- {line[2:]}")
        # Detect numbered lists
        elif re.match(r"^\d+[\.\)]\s", line):
            markdown_parts.append(line)
        else:
            markdown_parts.append(line)
    
    return "\n".join(markdown_parts)


def convert_to_markdown(content: bytes, filename: str) -> str:

    extension = filename.lower().split(".")[-1]
    
    if extension == "pdf":
        return pdf_to_markdown(content)
    elif extension in ["docx", "doc"]:
        return docx_to_markdown(content)
    elif extension == "txt":
        return txt_to_markdown(content)
    elif extension == "md":
        # Already markdown, just decode
        return content.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {extension}")


def get_supported_extensions() -> list:
    """Get list of supported file extensions"""
    return ["pdf", "docx", "doc", "txt", "md"]
